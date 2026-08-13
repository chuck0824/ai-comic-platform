from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[2]
SHOT = ROOT / "artifacts/aicp-product-introduction/screenshots"
OUT = ROOT / "artifacts/AICP产品介绍_客户与投资人版_2026-08-14.docx"

NAVY = "0B1739"
BLUE = "2357E6"
CYAN = "15B8A6"
PURPLE = "7657F6"
INK = "172033"
MUTED = "5D667A"
LIGHT = "F2F5FA"
PALE_BLUE = "EAF0FF"
PALE_CYAN = "E8F8F5"
PALE_PURPLE = "F0EDFF"
WHITE = "FFFFFF"
LINE = "DCE3EE"
FONT = "Noto Sans CJK SC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size=8):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_row_height(row, twips, exact=False):
    tr_pr = row._tr.get_or_add_trPr()
    tr_h = OxmlElement("w:trHeight")
    tr_h.set(qn("w:val"), str(twips))
    tr_h.set(qn("w:hRule"), "exact" if exact else "atLeast")
    tr_pr.append(tr_h)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def add_text(p, text, size=10.5, color=INK, bold=False, name=FONT):
    r = p.add_run(text)
    set_run_font(r, name)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return r


def clear_para(p):
    for r in list(p.runs):
        p._element.remove(r._element)


def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.45)
    sec.left_margin = Cm(1.65)
    sec.right_margin = Cm(1.65)
    sec.header_distance = Cm(0.65)
    sec.footer_distance = Cm(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 8),
        ("Heading 1", 21, NAVY, 0, 8),
        ("Heading 2", 14, BLUE, 10, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        st = doc.styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Kicker" not in [s.name for s in doc.styles]:
        st = doc.styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(8.5)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE)
        st.paragraph_format.space_after = Pt(3)


def add_page_number(p):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, "AICP 产品介绍  ·  ", 8, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def add_header_footer(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    add_text(hp, "AICP  |  AI CONTENT PRODUCTION", 7.5, MUTED, True)
    ppr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "3")
    bot.set(qn("w:color"), LINE)
    pbdr.append(bot)
    ppr.append(pbdr)
    add_page_number(sec.footer.paragraphs[0])
    clear_para(sec.first_page_header.paragraphs[0])
    clear_para(sec.first_page_footer.paragraphs[0])


def page_break(doc):
    doc.add_page_break()


def page_title(doc, kicker, title, intro=None):
    p = doc.add_paragraph(style="Kicker")
    add_text(p, kicker.upper(), 8.5, BLUE, True)
    p = doc.add_paragraph(style="Heading 1")
    add_text(p, title, 21, NAVY, True)
    if intro:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(9)
        add_text(p, intro, 10.5, MUTED)


def add_rule(doc, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(8)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "16")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pbdr.append(bot)
    ppr.append(pbdr)


def add_cards(doc, cards, cols=3, fills=None):
    rows = (len(cards) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = Inches(6.9 / cols)
    for i, (title, body) in enumerate(cards):
        cell = table.cell(i // cols, i % cols)
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, 150, 170, 150, 170)
        set_cell_border(cell, WHITE, 10)
        set_cell_shading(cell, (fills or [PALE_BLUE, PALE_CYAN, PALE_PURPLE])[i % len(fills or [1, 2, 3])])
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        add_text(p, title, 11, NAVY, True)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_text(p, body, 8.8, MUTED)
    for r in table.rows:
        set_row_height(r, 1150)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc, items, color=BLUE):
    for title, body in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.1)
        p.paragraph_format.first_line_indent = Cm(-0.1)
        p.paragraph_format.space_after = Pt(6)
        add_text(p, "●  ", 8.5, color, True)
        add_text(p, title + "：", 10.5, NAVY, True)
        add_text(p, body, 10.2, INK)


def add_callout(doc, title, body, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 150, 200, 150, 200)
    set_cell_border(cell, accent, 10)
    p = cell.paragraphs[0]
    add_text(p, title, 10.5, NAVY, True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_text(p, body, 9.5, MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_screenshot(doc, filename, caption, value, width=6.85):
    path = SHOT / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_text(p, caption, 9.2, NAVY, True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_text(p, value, 8.8, MUTED)


def add_architecture(doc):
    t = doc.add_table(rows=3, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    widths = [2.85, 0.85, 2.85]
    for row in t.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 120, 140, 120, 140)
            set_cell_border(cell, WHITE, 8)
    data = [
        ("内容生产工作台 · 8080", "创作与生产", "模型服务控制台 · 3001"),
        ("剧本 · 分镜 · 画布 · 资产", "任务 / 结果", "模型 · 渠道 · 路由 · 计量"),
        ("Agent · SOP · 企业业务", "状态 / 成本", "日志 · 额度 · 可观测性"),
    ]
    for r, rowdata in enumerate(data):
        for c, txt in enumerate(rowdata):
            cell = t.cell(r, c)
            set_cell_shading(cell, NAVY if c != 1 else PALE_CYAN)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, txt, 10.5 if r == 0 else 9.2, WHITE if c != 1 else NAVY, r == 0 or c == 1)
        set_row_height(t.rows[r], 840 if r == 0 else 680)
    return t


def add_flow(doc):
    steps = ["故事种子", "剧本 / 世界观", "分集 / 分镜", "画布编排", "多模型生成", "Agent / SOP", "资产 / 交付"]
    t = doc.add_table(rows=2, cols=len(steps))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, step in enumerate(steps):
        cell = t.cell(0, i)
        set_cell_shading(cell, [NAVY, BLUE, PURPLE, NAVY, BLUE, PURPLE, CYAN][i])
        set_cell_border(cell, WHITE, 8)
        set_cell_margins(cell, 120, 70, 120, 70)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"{i+1:02d}", 8, WHITE, True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_text(p, step, 8.5, WHITE, True)
        arrow = t.cell(1, i)
        set_cell_border(arrow, WHITE, 0)
        p = arrow.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, "→" if i < len(steps)-1 else "✓", 13, CYAN if i == len(steps)-1 else BLUE, True)
    return t


doc = Document()
style_doc(doc)
add_header_footer(doc)

# 1 Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(28)
add_text(p, "AICP", 12, BLUE, True)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(4)
add_text(p, "AI 漫剧与视频内容", 30, NAVY, True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
add_text(p, "工业化生产工作台", 30, BLUE, True)
add_rule(doc, CYAN)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(22)
add_text(p, "从模型能力到可交付内容，让创作团队像一条可治理、可复用、可扩展的产线运行。", 15, INK, True)
add_cards(doc, [
    ("CONTENT", "剧本、分镜、画布与资产"),
    ("INTELLIGENCE", "Agent、模型与工作流"),
    ("GOVERNANCE", "质量、成本与企业治理"),
], 3)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(54)
add_text(p, "客户与投资人介绍版", 11, NAVY, True)
p = doc.add_paragraph()
add_text(p, "2026 年 8 月  ·  PRODUCT BRIEF", 9, MUTED)

# 2 Summary
page_break(doc)
page_title(doc, "01 · Executive Summary", "一套平台，连接创意、模型与交付", "AICP 的核心不是多一个生成入口，而是把内容生产全过程变成可管理的系统。")
add_cards(doc, [
    ("面向谁", "漫剧、短剧、长篇内容与 TVC 团队；创作者、制片、企业管理者与模型运营人员。"),
    ("解决什么", "工具割裂、过程失控、经验难复制、模型接入重复以及成本与交付难追踪。"),
    ("如何实现", "内容生产工作台负责业务流程，模型服务控制台负责模型、渠道、计量与可观测性。"),
], 3)
add_callout(doc, "对客户的价值", "减少跨工具搬运，让项目、任务、版本、资产、质量和成本进入同一条可追踪生产链。", PALE_CYAN, CYAN)
add_callout(doc, "对投资人的意义", "从一次性生成工具升级为深入业务流程的生产基础设施，价值可由席位、用量、企业治理与生态服务多层承载。", PALE_PURPLE, PURPLE)
add_bullets(doc, [
    ("正确边界", "定位为生产工作台，不提供视频剪辑或多轨时间轴。"),
    ("事实口径", "本文使用真实运行页面；空状态只证明产品结构，不代表经营规模。"),
])

# 3 Challenges
page_break(doc)
page_title(doc, "02 · Market Problem", "内容生产进入 AI 时代，管理方式仍停留在工具时代")
add_cards(doc, [
    ("01 工具割裂", "文本、图像、视频与配音分散，项目上下文反复搬运。"),
    ("02 过程不可控", "生成结果很多，采用版本、失败原因和交付状态难追踪。"),
    ("03 经验难复制", "导演、编剧与制片判断停留在个人，难固化为组织流程。"),
    ("04 接入重复", "多供应商、多协议、多计费口径增加研发和运维负担。"),
    ("05 成本不透明", "难回答谁调用了什么模型、花了多少、是否产生有效结果。"),
    ("06 治理缺位", "权限、预算、审批、审计和资产边界无法覆盖 AI 生产链。"),
], 3)
add_callout(doc, "关键判断", "行业缺少的不是又一个模型入口，而是一套能够把创意、模型、流程、人和资产连接起来的生产操作系统。", NAVY, CYAN)

# 4 Positioning & architecture
page_break(doc)
page_title(doc, "03 · Product Positioning", "双平台、同一条生产链", "内容体验与模型底座分层建设，既服务创作者，也服务技术与运营团队。")
add_architecture(doc)
doc.add_paragraph()
add_bullets(doc, [
    ("8080 内容生产工作台", "承载内容项目、剧本、分镜、画布、资产、Agent、SOP 与企业业务。"),
    ("3001 模型服务控制台", "承载模型目录、上游渠道、路由、额度、日志和调用成本。"),
    ("协同关系", "生产端提出任务，模型端统一调用和计量，结果回到资产与任务体系沉淀。"),
])
add_callout(doc, "当前账号边界", "两端通过短时 SSO/JWT 桥与用户映射协同；完整统一账户仍处于演进过程中，本文不将目标态表述为已完成。", PALE_BLUE, BLUE)

# 5 Flow
page_break(doc)
page_title(doc, "04 · Production Loop", "从故事种子，到可交付素材包", "每个生成结果都进入版本、任务、质量、成本和资产关系，而不是停留在聊天窗口。")
add_flow(doc)
doc.add_paragraph()
add_cards(doc, [
    ("可追踪", "知道结果从哪里来、用过什么模型、经历过哪些审核。"),
    ("可复用", "角色、场景、镜头、提示与工作流持续沉淀为项目资产。"),
    ("可治理", "生产准入、返工、权限、预算与审计覆盖关键动作。"),
], 3)
add_callout(doc, "闭环价值", "AICP 优化的不是一次提示词，而是从创作决策到素材交付的整条价值链。", NAVY, CYAN)

# 6 Home
page_break(doc)
page_title(doc, "05 · Unified Workspace", "三类创作，共用一套生产底座")
add_screenshot(doc, "01-8080-home.png", "图 1｜AICP 内容生产工作台首页", "短剧、长篇与 TVC 从同一入口创建；创作、画布、内容交易、Agent、资产、企业与 SOP 在统一导航中协同。")
add_callout(doc, "业务价值", "内容类型可以不同，但项目、任务、资产和治理方式保持一致，降低团队扩大品类时的系统复杂度。", PALE_CYAN, CYAN)

# 7 Scripts
page_break(doc)
page_title(doc, "06 · Script & IP Assets", "剧本不是文件，而是生产链的上游资产")
add_screenshot(doc, "02-8080-script-gen.png", "图 2｜剧本创作启动台", "快速创作、专业分步创作、已有文稿导入与 TVC 创作覆盖不同成熟度和内容类型。", 5.75)
add_screenshot(doc, "03-8080-warehouse.png", "图 3｜剧本仓库", "草稿、审核、锁稿、生产、完成与归档状态连接内容决策和后续生产。", 5.75)

# 8 Canvas
page_break(doc)
page_title(doc, "07 · Visual Production", "用画布编排生产，而不是堆叠工具")
add_screenshot(doc, "04-8080-canvas-projects.png", "图 4｜画布项目中心", "画布承接剧本与分镜之后的可视化生产，通过项目、状态与模式组织生产单元。")
add_bullets(doc, [
    ("节点化", "角色、场景、分镜、图片、视频和任务可以在同一空间组织。"),
    ("批量化", "工作流模板和生成任务让重复步骤可复用、可规模化。"),
    ("边界清晰", "平台聚焦生产编排与素材交付，不替代专业剪辑器。"),
])

# 9 Assets
page_break(doc)
page_title(doc, "08 · Asset Operations", "成功结果和失败任务，都进入同一个资产视图")
add_screenshot(doc, "05-8080-asset-workbench.png", "图 5｜资产生成工作台", "页面展示排队中、失败、任务类型、模型与错误原因，异常不再散落在各个供应商后台。")
add_cards(doc, [
    ("生成记录", "保留任务状态、模型和执行结果。"),
    ("资产组织", "按项目、集合与业务类型分类管理。"),
    ("生产复用", "资产可发送到画布、重新生成、归档或进入后续交付。"),
], 3)

# 10 Agent/SOP
page_break(doc)
page_title(doc, "09 · Agent & SOP", "把个人经验，变成可配置的组织能力")
add_screenshot(doc, "06-8080-agent-config.png", "图 6｜Agent 配置中心", "Agent 拥有独立配置入口，为角色蓝图、版本、绑定、试跑和执行快照提供承载。", 5.75)
add_screenshot(doc, "07-8080-sop.png", "图 7｜工业化生产 SOP", "以项目为单位查看生产准入与返工工单；权威规格进一步定义生产 Gate、规则引擎和失败恢复。", 5.75)

# 11 Enterprise
page_break(doc)
page_title(doc, "10 · Enterprise Governance", "让一支团队，像一个系统运转", "企业价值不止来自更多席位，而来自生产动作进入权限、预算、审批和审计体系。")
add_cards(doc, [
    ("Workspace", "个人空间与企业空间形成清晰的数据和资产边界。"),
    ("组织权限", "部门、成员、角色与数据范围决定谁能看、谁能操作。"),
    ("统一审批", "采购、资产发布和项目导出等关键动作集中呈现。"),
    ("采购预算", "预算是治理配额，账户余额与账本仍由模型服务账户域负责。"),
    ("跨域审计", "内容、资产、交易和交付的关键操作形成可追溯记录。"),
    ("故障关闭", "依赖账户域的关键操作在不可用时拒绝伪成功，保护权限与资金安全。"),
], 3)
add_callout(doc, "当前状态说明", "企业能力采用领域化渐进完成方式；本地企业概览接口仍需继续联调，因此本页只呈现已确认的架构与产品边界，不使用异常页面截图。", PALE_PURPLE, PURPLE)

# 12 Models foundation
page_break(doc)
page_title(doc, "11 · Model Infrastructure", "把多模型复杂度，收敛为一个可管理的底座")
add_screenshot(doc, "10-3001-channels.png", "图 8｜模型渠道管理", "统一管理上游供应商、渠道状态、类型、筛选、路由与重试设置。", 5.75)
add_screenshot(doc, "11-3001-models.png", "图 9｜模型元数据管理", "模型名称、供应商、端点、标签、授权包与部署信息进入统一目录。", 5.75)

# 13 Observability
page_break(doc)
page_title(doc, "12 · Cost & Performance", "调用、性能与成本，在同一张图上被看见")
add_screenshot(doc, "08-3001-model-dashboard.png", "图 10｜模型调用分析", "统一观察调用量、额度、Token、RPM、TPM、消耗分布和趋势；空数据为当前环境状态，不代表经营结果。")
add_callout(doc, "管理价值", "模型能力只有进入用量、性能和成本的共同视图，才真正具备企业级可运营性。", PALE_CYAN, CYAN)

# 14 Logs
page_break(doc)
page_title(doc, "13 · Traceability", "同步调用与异步任务，采用两套可追踪视图")
add_screenshot(doc, "12-3001-usage-logs.png", "图 11｜使用日志", "按时间、渠道、用户、令牌、模型、Token 与费用追踪同步请求。", 5.75)
add_screenshot(doc, "13-3001-task-logs.png", "图 12｜任务日志", "异步生成任务进一步记录任务编号、运行时间、状态与详情。", 5.75)

# 15 Moat
page_break(doc)
page_title(doc, "14 · Differentiation", "AICP 的壁垒，不是某一个模型")
add_cards(doc, [
    ("生产闭环", "从创意到交付，价值跨越多个角色和生产阶段。"),
    ("双平台协同", "业务体验与模型底座分层，保持独立演进和清晰责任。"),
    ("可治理 AI", "Agent、SOP、审批、预算和审计将生成纳入企业规则。"),
    ("资产关系", "版本、任务、提示、结果与采用关系持续形成复用网络。"),
    ("模型适配", "统一渠道、协议和元数据，降低模型供给变化对业务的冲击。"),
    ("流程数据", "质量、失败、返工、成本和交付数据积累为持续优化的基础。"),
], 3)
add_callout(doc, "飞轮逻辑", "模型越丰富，统一编排和治理越有价值；项目使用越深入，流程与资产关系越完整，迁移成本也越高。", NAVY, CYAN)

# 16 Customer value
page_break(doc)
page_title(doc, "15 · Customer Value", "客户购买的不是功能清单，而是确定性")
add_cards(doc, [
    ("效率", "统一入口减少跨工具切换、重复录入和上下文搬运。"),
    ("质量", "分镜、版本、生产 Gate 与返工机制让审核有依据。"),
    ("成本", "模型路由、用量和日志让费用可见、可定位。"),
    ("协作", "项目、Workspace、角色与任务形成团队共同事实源。"),
    ("资产", "角色、场景、镜头和生成记录从个人文件变成可复用资产。"),
    ("交付", "采用版本、质量状态和素材清单可追踪，降低交接风险。"),
], 3)
add_bullets(doc, [
    ("内容制作公司", "提升多项目、多团队并行生产的一致性和可控性。"),
    ("MCN / 漫剧团队", "用标准流程降低新人上手与大批量生产的管理成本。"),
    ("品牌与企业客户", "让生成内容进入权限、预算、审计和合规边界。"),
])

# 17 Investment
page_break(doc)
page_title(doc, "16 · Business & Investment", "从创作者工具，向内容生产基础设施延展")
add_cards(doc, [
    ("席位与团队版", "以创作者和团队协作为入口，建立稳定使用关系。"),
    ("模型用量", "通过统一模型调用与增值服务承接生产规模增长。"),
    ("企业治理", "权限、预算、审批、审计和私有化形成更高价值方案。"),
    ("开放集成", "API、模型渠道与业务系统集成扩大平台边界。"),
    ("内容与资产服务", "交易、授权和生态服务提供进一步商业化空间。"),
    ("联合解决方案", "与内容机构、模型厂商和交付伙伴共同沉淀行业模板。"),
], 3)
add_callout(doc, "规模化路径", "创作者工具 → 团队生产系统 → 企业控制面 → 内容与资产生态。本文不提供未经验证的收入预测或市场规模数字。", PALE_PURPLE, PURPLE)

# 18 Close
page_break(doc)
page_title(doc, "17 · Partnership", "从模型到作品，从工具到基础设施")
add_cards(doc, [
    ("团队 / 企业工作台", "围绕客户内容类型和生产 SOP 配置工作空间与流程。"),
    ("私有化部署", "结合数据边界、合规要求和模型供应设计部署方案。"),
    ("模型与渠道接入", "统一上游模型、路由、调用观察与成本治理。"),
    ("API 与系统集成", "连接客户已有内容、资产、审批和交付系统。"),
    ("流程共建", "将编剧、分镜、导演、质检与交付经验沉淀为模板。"),
    ("联合解决方案", "与内容机构、模型厂商和行业伙伴共同打造场景方案。"),
], 3)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p, "AICP", 26, BLUE, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p, "让每一次生成，都成为可管理、可复用、可交付的生产资产。", 14, NAVY, True)
add_rule(doc, CYAN)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
add_text(p, "说明：本文依据 2026-08-14 本地运行页面与 docs/01-core/ 权威产品资料整理。页面空状态只用于展示产品结构，不代表客户规模或经营成果。", 8.5, MUTED)


# Document properties
doc.core_properties.title = "AICP 产品介绍｜客户与投资人版"
doc.core_properties.subject = "AI 漫剧与视频内容工业化生产工作台"
doc.core_properties.author = "AICP"
doc.core_properties.keywords = "AICP, AI漫剧, 内容生产, 模型网关, Agent, SOP"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
